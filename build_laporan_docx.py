from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_placeholder_box(doc, label="Tempel screenshot di sini"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.2)
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n[ " + label + " ]\n")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(120, 120, 120)
    p2 = cell.add_paragraph("Gunakan screenshot yang jelas, menampilkan method, URL, request body/params, header auth bila diperlukan, dan response.")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p2.runs:
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(110, 110, 110)
    doc.add_paragraph()


def add_explanation_block(doc, explanation, objective=None):
    p = doc.add_paragraph()
    p.style = "List Bullet"
    run = p.add_run("Tujuan pengujian: ")
    run.bold = True
    p.add_run(objective or "Memastikan endpoint berjalan sesuai fungsi yang diharapkan.")

    p2 = doc.add_paragraph()
    p2.style = "List Bullet"
    run = p2.add_run("Penjelasan screenshot: ")
    run.bold = True
    p2.add_run(explanation)

    p3 = doc.add_paragraph()
    p3.style = "List Bullet"
    run = p3.add_run("Hasil yang diharapkan: ")
    run.bold = True
    p3.add_run("Response berhasil muncul dan membuktikan endpoint dapat dipanggil dengan benar.")

    doc.add_paragraph()


def add_shot(doc, number, title, explanation, objective=None):
    doc.add_heading(f"{number}. {title}", level=2)
    add_placeholder_box(doc, f"Screenshot {number}: {title}")
    add_explanation_block(doc, explanation, objective)


doc = Document()

for section in doc.sections:
    set_page_margins(section)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(11)
styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(22)
styles["Heading 1"].font.name = "Arial"
styles["Heading 1"].font.size = Pt(15)
styles["Heading 2"].font.name = "Arial"
styles["Heading 2"].font.size = Pt(12)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAPORAN UJIAN AKHIR SEMESTER\nWEB SERVICE APPLICATION")
r.bold = True
r.font.size = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\nPengembangan REST API Mini E-Commerce\nmenggunakan Laravel, Sanctum, dan Scribe")
r.font.size = Pt(14)

doc.add_paragraph()

meta = doc.add_table(rows=5, cols=2)
meta.style = "Table Grid"
labels = ["Nama", "NPM", "Kelas", "Mata Kuliah", "Tanggal Penyusunan"]
for i, label in enumerate(labels):
    meta.cell(i, 0).text = label
    meta.cell(i, 1).text = ": ............................................................"
doc.add_paragraph()

p = doc.add_paragraph(
    "Catatan: file ini disusun agar screenshot pengujian dapat ditempel langsung pada bagian yang sudah disediakan. "
    "Setiap screenshot telah dilengkapi penjelasan singkat sehingga laporan mudah dibaca dan sesuai kebutuhan pengumpulan UAS."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_page_break()

doc.add_heading("A. Pendahuluan", level=1)
for text in [
    "Pada proyek UAS ini dikembangkan sebuah REST API mini e-commerce menggunakan framework Laravel. API ini mencakup autentikasi berbasis token menggunakan Laravel Sanctum, pengelolaan kategori produk, pengelolaan produk, pengelolaan transaksi, serta dokumentasi API interaktif menggunakan Scribe.",
    "Tujuan penyusunan laporan ini adalah mendokumentasikan hasil pengujian seluruh endpoint API yang telah dibuat. Setiap pengujian dilakukan menggunakan Postman, kemudian hasil respons dari masing-masing endpoint dibuktikan melalui screenshot yang nantinya ditempelkan pada tempat yang telah disediakan di dokumen ini.",
]:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_heading("B. Ringkasan Fitur API", level=1)
for item in [
    "Autentikasi user: register, login, dan logout dengan token Sanctum.",
    "Manajemen kategori: menampilkan, menambah, mengubah, dan menghapus kategori.",
    "Manajemen produk: CRUD produk, upload gambar, pagination, search, dan relasi kategori.",
    "Manajemen transaksi: CRUD transaksi yang terhubung dengan user dan produk.",
    "Dokumentasi API otomatis menggunakan Scribe pada endpoint /docs.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("C. Petunjuk Pengisian Screenshot", level=1)
for item in [
    "Tempel screenshot tepat di dalam kotak placeholder yang tersedia.",
    "Usahakan setiap screenshot memperlihatkan method, URL endpoint, request body atau query parameter, serta response JSON.",
    "Untuk endpoint yang membutuhkan autentikasi, pastikan header Bearer Token terlihat atau setidaknya request berhasil dibuktikan oleh response.",
    "Jika satu screenshot dirasa terlalu padat, dapat diganti dengan dua screenshot berurutan pada bagian yang sama.",
]:
    doc.add_paragraph(item, style="List Number")

doc.add_page_break()

doc.add_heading("D. Hasil Pengujian Endpoint API", level=1)

shots = [
    (
        "1",
        "POST /api/auth/register",
        "Screenshot ini menunjukkan proses pendaftaran user baru melalui endpoint register. "
        "Request berisi nama, email, password, dan konfirmasi password. Response yang berhasil menandakan data user baru tersimpan dan sistem mengembalikan token autentikasi.",
        "Memastikan fitur registrasi user berjalan dengan baik.",
    ),
    (
        "2",
        "POST /api/auth/login",
        "Screenshot ini menunjukkan proses login menggunakan akun yang telah tersedia. "
        "Response yang berhasil akan menampilkan token yang nantinya digunakan untuk mengakses endpoint yang dilindungi middleware auth:sanctum.",
        "Memastikan user dapat login dan memperoleh token akses.",
    ),
    (
        "3",
        "POST /api/auth/logout",
        "Screenshot ini menunjukkan proses logout dari user yang sedang aktif. "
        "Request dilakukan menggunakan Bearer Token yang sebelumnya diperoleh saat login. Response sukses menandakan token yang digunakan telah dihapus sehingga sesi akses API berakhir.",
        "Memastikan fitur logout dan penghapusan token berhasil dilakukan.",
    ),
    (
        "4",
        "GET /api/categories",
        "Screenshot ini menunjukkan hasil pengambilan daftar kategori. "
        "Response menampilkan data kategori dalam bentuk pagination sehingga membuktikan endpoint kategori dapat diakses setelah autentikasi berhasil.",
        "Memastikan endpoint daftar kategori berjalan dan mengembalikan data dengan benar.",
    ),
    (
        "5",
        "POST /api/categories",
        "Screenshot ini menunjukkan proses penambahan kategori baru. "
        "Request berisi nama kategori dan deskripsi, sedangkan response menampilkan data kategori yang baru saja berhasil dibuat.",
        "Memastikan fitur tambah kategori dapat menyimpan data ke database.",
    ),
    (
        "6",
        "PUT /api/categories/{id}",
        "Screenshot ini menunjukkan proses pembaruan data kategori berdasarkan id kategori yang dipilih. "
        "Response berhasil menandakan data kategori telah berubah sesuai input terbaru.",
        "Memastikan fitur update kategori berjalan dengan benar.",
    ),
    (
        "7",
        "DELETE /api/categories/{id}",
        "Screenshot ini menunjukkan proses penghapusan kategori. "
        "Response sukses menandakan kategori yang dipilih telah berhasil dihapus dari sistem.",
        "Memastikan fitur hapus kategori bekerja sesuai harapan.",
    ),
    (
        "8",
        "GET /api/products",
        "Screenshot ini menunjukkan daftar produk yang ditampilkan dengan relasi kategori. "
        "Bagian response harus memperlihatkan bahwa setiap produk memiliki data kategori secara nested JSON.",
        "Memastikan endpoint daftar produk menampilkan relasi kategori (eager loading).",
    ),
    (
        "9",
        "GET /api/products?search=...&per_page=...",
        "Screenshot ini menunjukkan fitur pencarian dan pagination pada endpoint produk. "
        "Response membuktikan bahwa sistem dapat memfilter data berdasarkan kata kunci tertentu serta membatasi jumlah data per halaman.",
        "Memastikan fitur search dan pagination produk berjalan.",
    ),
    (
        "10",
        "GET /api/products/{id}",
        "Screenshot ini menunjukkan detail satu produk tertentu berdasarkan id produk. "
        "Response harus menampilkan data lengkap produk, termasuk kategori yang terkait.",
        "Memastikan endpoint detail produk berfungsi dengan benar.",
    ),
    (
        "11",
        "POST /api/products",
        "Screenshot ini menunjukkan proses penambahan produk baru beserta upload gambar. "
        "Pada request biasanya digunakan form-data agar file gambar dapat dikirim, lalu response menampilkan produk yang baru berhasil dibuat.",
        "Memastikan fitur tambah produk dan upload gambar berjalan.",
    ),
    (
        "12",
        "PUT /api/products/{id}",
        "Screenshot ini menunjukkan proses pembaruan produk, baik perubahan data teks maupun penggantian gambar produk. "
        "Response sukses membuktikan data produk berhasil diperbarui.",
        "Memastikan fitur update produk bekerja dengan baik.",
    ),
    (
        "13",
        "DELETE /api/products/{id}",
        "Screenshot ini menunjukkan proses penghapusan produk. "
        "Response yang berhasil menandakan produk dan file gambar terkait telah dihapus sesuai logika sistem.",
        "Memastikan fitur hapus produk berjalan dengan benar.",
    ),
    (
        "14",
        "GET /api/transactions",
        "Screenshot ini menunjukkan daftar transaksi yang ditampilkan dengan pagination. "
        "Response harus memperlihatkan bahwa transaksi terhubung dengan data user dan produk.",
        "Memastikan endpoint daftar transaksi berjalan dan relasinya terbaca.",
    ),
    (
        "15",
        "GET /api/transactions?status=...&per_page=...",
        "Screenshot ini menunjukkan pengujian filter status dan pagination pada transaksi. "
        "Response membuktikan bahwa transaksi dapat difilter, misalnya berdasarkan status paid, pending, atau cancelled.",
        "Memastikan fitur filter transaksi berfungsi.",
    ),
    (
        "16",
        "GET /api/transactions/{id}",
        "Screenshot ini menunjukkan detail satu transaksi berdasarkan id. "
        "Response menampilkan informasi transaksi secara lengkap, termasuk user yang melakukan transaksi dan produk yang dibeli.",
        "Memastikan endpoint detail transaksi berjalan dengan benar.",
    ),
    (
        "17",
        "POST /api/transactions",
        "Screenshot ini menunjukkan proses pembuatan transaksi baru. "
        "Request biasanya berisi user_id, product_id, quantity, dan status. Response akan menampilkan hasil perhitungan harga satuan dan total harga secara otomatis.",
        "Memastikan fitur tambah transaksi berjalan sesuai logika bisnis.",
    ),
    (
        "18",
        "PUT /api/transactions/{id}",
        "Screenshot ini menunjukkan proses pembaruan transaksi, misalnya perubahan quantity atau status transaksi. "
        "Response sukses membuktikan bahwa data transaksi berhasil diperbarui.",
        "Memastikan fitur update transaksi bekerja dengan baik.",
    ),
    (
        "19",
        "DELETE /api/transactions/{id}",
        "Screenshot ini menunjukkan proses penghapusan transaksi. "
        "Response berhasil menandakan transaksi yang dipilih sudah terhapus dari sistem.",
        "Memastikan fitur hapus transaksi berjalan dengan benar.",
    ),
    (
        "20",
        "Halaman Dokumentasi Scribe (/docs)",
        "Screenshot ini menunjukkan dokumentasi API yang dihasilkan otomatis oleh Scribe. "
        "Halaman ini membuktikan bahwa dokumentasi endpoint telah berhasil dibuat dan dapat diakses melalui browser secara interaktif.",
        "Memastikan dokumentasi API tersedia sesuai instruksi UAS.",
    ),
]

for number, title, explanation, objective in shots:
    add_shot(doc, number, title, explanation, objective)

doc.add_heading("E. Kesimpulan", level=1)
for text in [
    "Berdasarkan hasil pengujian yang telah dilakukan, REST API mini e-commerce yang dibangun telah mencakup fitur utama sesuai instruksi UAS, yaitu autentikasi dengan Sanctum, relasi kategori dan produk, pagination dan search, CRUD produk dengan upload gambar, tabel transaksi sebagai nilai tambah, serta dokumentasi otomatis menggunakan Scribe.",
    "Melalui pengujian menggunakan Postman dan dokumentasi visual melalui screenshot, dapat disimpulkan bahwa setiap endpoint API telah berhasil diimplementasikan dan dapat berjalan sesuai fungsinya masing-masing.",
]:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

output = "/Users/pesal/Documents/UAS_WSA/Laporan_UAS_WSA_Template.docx"
doc.save(output)
print(output)
